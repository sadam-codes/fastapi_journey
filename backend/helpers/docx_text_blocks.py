"""Read/write plain text per python-docx paragraph in the same order as merge (body, then table cells)."""

import io
from typing import Any

from fastapi import HTTPException, status

MAX_BLOCKS = 4000
MAX_TOTAL_CHARS = 2_000_000


def _paragraph_plain(p: Any) -> str:
    return "".join(r.text for r in p.runs)


def _set_paragraph_plain(p: Any, text: str) -> None:
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def extract_docx_text_blocks(blob: bytes) -> list[str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not available.",
        ) from exc

    doc = Document(io.BytesIO(blob))
    blocks: list[str] = []
    for p in doc.paragraphs:
        blocks.append(_paragraph_plain(p))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    blocks.append(_paragraph_plain(p))

    if len(blocks) > MAX_BLOCKS:
        raise HTTPException(
            status_code=400,
            detail=f"Document has too many text blocks ({len(blocks)}). Max supported is {MAX_BLOCKS}.",
        )
    total = sum(len(b) for b in blocks)
    if total > MAX_TOTAL_CHARS:
        raise HTTPException(
            status_code=400,
            detail="Document text is too large to edit in the browser.",
        )
    return blocks


def apply_docx_text_blocks(blob: bytes, blocks: list[str]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not available.",
        ) from exc

    if len(blocks) > MAX_BLOCKS:
        raise HTTPException(status_code=400, detail="Too many text blocks.")
    total = sum(len(b) for b in blocks)
    if total > MAX_TOTAL_CHARS:
        raise HTTPException(status_code=400, detail="Text blocks are too large.")

    doc = Document(io.BytesIO(blob))
    idx = 0

    def apply_paragraphs(paragraphs: Any) -> None:
        nonlocal idx
        for p in paragraphs:
            if idx >= len(blocks):
                raise HTTPException(
                    status_code=400,
                    detail="Too few text blocks: the document has more paragraphs/cells than the list you sent.",
                )
            _set_paragraph_plain(p, blocks[idx])
            idx += 1

    apply_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_paragraphs(cell.paragraphs)

    if idx != len(blocks):
        raise HTTPException(
            status_code=400,
            detail="Too many text blocks: the document has fewer paragraphs/cells than the list you sent.",
        )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
