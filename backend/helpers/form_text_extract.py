import io
import zipfile
from typing import Final

from fastapi import HTTPException, status
from PIL import Image
from pypdf import PdfReader

_ALLOWED_SUFFIXES: Final = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp", ".gif", ".tiff", ".bmp"}


def assert_allowed_form_filename(filename: str) -> str:
    lower = filename.lower().strip()
    if not any(lower.endswith(s) for s in _ALLOWED_SUFFIXES):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type. Allowed: {', '.join(sorted(_ALLOWED_SUFFIXES))}",
        )
    return lower


def extract_text_from_pdf(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def extract_text_from_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not available on the server.",
        ) from exc

    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n".join(parts).strip()


def extract_text_from_image(raw: bytes) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Image OCR is not available (pytesseract).",
        ) from exc

    try:
        img = Image.open(io.BytesIO(raw))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        return (pytesseract.image_to_string(img) or "").strip()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                "Could not read text from the image. Install Tesseract OCR on the server "
                "and ensure it is on PATH, or use PDF/DOCX with embedded text."
            ),
        ) from exc


def extract_plain_text_from_upload(*, filename: str, raw: bytes) -> str:
    lower = assert_allowed_form_filename(filename)
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(raw)
    if lower.endswith(".docx"):
        return extract_text_from_docx(raw)
    if lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".tiff", ".bmp")):
        return extract_text_from_image(raw)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type.")
