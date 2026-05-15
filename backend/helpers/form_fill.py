import base64
import io
import re
from collections import defaultdict
from typing import Any

from fastapi import HTTPException, status
from PIL import Image, ImageDraw, ImageFont


def _decode_data_url_image(val: str) -> bytes | None:
    """Return raw image bytes from a data URL, or None."""
    if not isinstance(val, str) or not val.startswith("data:image/") or ";base64," not in val:
        return None
    try:
        b64 = val.split(",", 1)[1]
        return base64.b64decode(b64, validate=True)
    except Exception:
        return None


def _image_bytes_for_embed(val: str) -> bytes | None:
    """Decode data URL and normalize to PNG for reliable embedding in DOCX/PDF."""
    raw = _decode_data_url_image(val)
    if not raw:
        return None
    try:
        im = Image.open(io.BytesIO(raw))
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGBA")
        buf = io.BytesIO()
        im.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _signature_fallback_text(val: str) -> str:
    """Short line when no drawable signature image is available."""
    if val and val.strip() and not val.startswith("data:image/"):
        return str(val).strip()
    return "[Signed electronically]" if val and val.strip() else ""


def _checkbox_merge_replacement(raw: Any) -> str:
    """Checked/unchecked as ballot symbols — avoids printing true/false/yes/no in documents."""
    if raw is True:
        checked = True
    elif raw is False or raw is None:
        checked = False
    else:
        s = str(raw).strip().lower()
        checked = s in ("true", "1", "yes", "on")
    # ☑ checked; ☐ empty box for unchecked (multi-select leaves other boxes blank)
    return "\u2611" if checked else "\u2610"


# After merge, templates often still have literal ``(Yes)`` / ``(No)`` next to ``[[…]]`` → ``☑(Yes)``. Strip those parens.
_BALLOT_PAREN_LABEL_RX = re.compile(r"([\u2610\u2611])\s*\(([^)]{0,200})\)")


def _strip_parentheses_after_ballots(chunk: str) -> str:
    """``☑(Yes)`` → ``☑ Yes``; ``☐ (No)`` → ``☐ No``."""
    if not chunk:
        return chunk
    return _BALLOT_PAREN_LABEL_RX.sub(lambda m: f"{m.group(1)} {m.group(2).strip()}", chunk)


def _coalesce_string_segments_for_ballot_cleanup(segments: list[str | bytes]) -> list[str | bytes]:
    """Join adjacent string segments so ``☑`` + ``(Yes)`` can be cleaned; keep image bytes as boundaries."""
    out: list[str | bytes] = []
    buf: list[str] = []
    for seg in segments:
        if isinstance(seg, bytes):
            if buf:
                out.append(_strip_parentheses_after_ballots("".join(buf)))
                buf = []
            out.append(seg)
        else:
            buf.append(seg)
    if buf:
        out.append(_strip_parentheses_after_ballots("".join(buf)))
    return out


def _paragraph_plain_preserving_tabs(paragraph: Any) -> str:
    """Paragraph text including Word tab stops (``<w:tab/>``), not only ``run.text``.

    ``"".join(r.text for r in p.runs)`` drops tabs, which breaks alignment (e.g. right-side labels).
    """
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return "".join(r.text for r in paragraph.runs)

    parts: list[str] = []
    for run in paragraph.runs:
        for child in list(run._element):
            t = child.tag
            if t == qn("w:t"):
                parts.append(child.text or "")
            elif t == qn("w:tab"):
                parts.append("\t")
            elif t == qn("w:br"):
                parts.append("\n")

    out = "".join(parts)
    if out == "" and paragraph.text:
        return "".join(r.text for r in paragraph.runs)
    return out


def _spec_placeholders_match_index(row: dict[str, Any]) -> int | None:
    raw = row.get("placeholder_match_index")
    if raw is None:
        return None
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


# Strip explicit radio / checkbox scaffolding (not in schema placeholders).
_RADIO_MARKUP_START_RX = re.compile(r"\[\[\s*RADIO_START\s*:([\s\S]*?)\]\]", re.IGNORECASE)
_RADIO_MARKUP_END_RX = re.compile(r"\[\[\s*RADIO_END\s*(?:\]\]|\])", re.IGNORECASE)
_CHECKBOX_MARKUP_START_RX = re.compile(r"\[\[\s*CHECKBOX_START\s*:([\s\S]*?)\]\]", re.IGNORECASE)
_CHECKBOX_MARKUP_END_RX = re.compile(r"\[\[\s*CHECKBOX_END\s*(?:\]\]|\])", re.IGNORECASE)


def _explicit_radio_markup_strip_specs(full: str) -> list[tuple[str, str | bytes, int | None]]:
    """Literal substrings to remove ``[[RADIO_*]]`` / ``[[CHECKBOX_*]]`` markers from this paragraph."""
    seen: set[str] = set()
    out: list[tuple[str, str | bytes, int | None]] = []
    for rx in (
        _RADIO_MARKUP_START_RX,
        _RADIO_MARKUP_END_RX,
        _CHECKBOX_MARKUP_START_RX,
        _CHECKBOX_MARKUP_END_RX,
    ):
        for m in rx.finditer(full or ""):
            s = m.group(0)
            if s not in seen:
                seen.add(s)
                out.append((s, "", None))
    out.sort(key=lambda x: -len(x[0]))
    return out


def _explicit_radio_markup_strip_specs_from_pdf(doc: Any) -> list[tuple[str, str | bytes, int | None]]:
    """Collect unique explicit radio/checkbox markup literals across all pages."""
    seen: set[str] = set()
    out: list[tuple[str, str | bytes, int | None]] = []
    for page in doc:
        t = page.get_text() or ""
        for rx in (
            _RADIO_MARKUP_START_RX,
            _RADIO_MARKUP_END_RX,
            _CHECKBOX_MARKUP_START_RX,
            _CHECKBOX_MARKUP_END_RX,
        ):
            for m in rx.finditer(t):
                s = m.group(0)
                if s not in seen:
                    seen.add(s)
                    out.append((s, "", None))
    out.sort(key=lambda x: -len(x[0]))
    return out


def _radio_group_label_bracket_strip_specs(schema: list[dict[str, Any]]) -> list[tuple[str, str, int | None]]:
    """Replace ``[Caption]`` with ``Caption`` when it matches a radio row's ``group_label`` (clean merged .docx)."""
    seen: set[str] = set()
    out: list[tuple[str, str, int | None]] = []
    for row in schema or []:
        it = str(row.get("input_type") or "").lower()
        if it not in ("radio", "checkbox"):
            continue
        gl = str(row.get("group_label") or "").strip()
        if not gl or gl in seen:
            continue
        seen.add(gl)
        out.append((f"[{gl}]", gl, None))
    return out


def _placeholder_specs(
    schema: list[dict[str, Any]], answers: dict[str, Any]
) -> list[tuple[str, str | bytes, int | None]]:
    """Each entry is (placeholder, replacement, occurrence_index or None).

    ``occurrence_index`` is 0-based among identical placeholder strings in scan order
    (see ``detect_dynamic_fields``). ``None`` means every occurrence of ``placeholder``
    gets the same replacement (legacy rows without ``placeholder_match_index``).
    """
    out: list[tuple[str, str | bytes, int | None]] = []
    for row in schema:
        key = row["key"]
        it = str(row.get("input_type") or "text").strip().lower()
        rg = str(row.get("radio_group") or "").strip()
        if it == "radio" and rg:
            raw_ans = answers.get(rg)
        else:
            raw_ans = answers.get(key)
        val = str(raw_ans if raw_ans is not None else "")
        midx = _spec_placeholders_match_index(row)
        for ph in row.get("placeholders", []) or []:
            phs = str(ph)
            if not phs:
                continue
            if it == "signature":
                img = _image_bytes_for_embed(val)
                if img:
                    out.append((phs, img, midx))
                else:
                    out.append((phs, _signature_fallback_text(val), midx))
            elif it == "checkbox":
                out.append((phs, _checkbox_merge_replacement(raw_ans), midx))
            elif it == "radio":
                opt = str(row.get("radio_option") or key).strip()
                sel = val.strip()
                out.append(
                    (phs, _checkbox_merge_replacement(bool(sel) and sel.lower() == opt.lower()), midx)
                )
            else:
                out.append((phs, val, midx))
    out.sort(key=lambda x: -len(x[0]))
    return out


def _full_string_to_segments(
    full: str,
    specs: list[tuple[str, str | bytes, int | None]],
    occ_state: defaultdict[str, int],
) -> list[str | bytes]:
    """Split document text into alternating literal strings and raw PNG segments."""
    out: list[str | bytes] = []
    buf: list[str] = []
    i = 0
    n = len(full)

    def flush_buf() -> None:
        nonlocal buf
        if buf:
            out.append("".join(buf))
            buf = []

    while i < n:
        matched = False
        for ph, repl, midx in specs:
            if not ph or not full.startswith(ph, i):
                continue
            k = occ_state[ph]
            if midx is not None and midx != k:
                continue
            flush_buf()
            if isinstance(repl, bytes):
                out.append(repl)
            elif repl:
                out.append(repl)
            i += len(ph)
            occ_state[ph] = k + 1
            matched = True
            break
        if not matched:
            buf.append(full[i])
            i += 1
    flush_buf()
    return out


def _rebuild_docx_paragraph(
    paragraph,
    full: str,
    specs: list[tuple[str, str | bytes, int | None]],
    occ_state: defaultdict[str, int],
) -> None:
    try:
        from docx.oxml.ns import qn
        from docx.shared import Inches
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not available.",
        ) from exc

    segments = _coalesce_string_segments_for_ballot_cleanup(_full_string_to_segments(full, specs, occ_state))
    p_el = paragraph._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    for seg in segments:
        if isinstance(seg, bytes):
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(seg), width=Inches(1.65))
        else:
            paragraph.add_run(seg)


def _iter_docx_paragraphs(doc: Any):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                yield p
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def fill_docx(blob: bytes, schema: list[dict[str, Any]], answers: dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not available.",
        ) from exc

    doc = Document(io.BytesIO(blob))
    strip_specs = _radio_group_label_bracket_strip_specs(schema)
    main_specs = _placeholder_specs(schema, answers)
    base_specs = sorted(strip_specs + main_specs, key=lambda x: -len(x[0]))
    occ_state: defaultdict[str, int] = defaultdict(int)
    for p in _iter_docx_paragraphs(doc):
        full = _paragraph_plain_preserving_tabs(p)
        radio_strip = _explicit_radio_markup_strip_specs(full)
        specs = sorted(radio_strip + base_specs, key=lambda x: -len(x[0]))
        if not full or not any(ph in full for ph, _, __ in specs):
            continue
        _rebuild_docx_paragraph(p, full, specs, occ_state)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _pdf_collect_hits(doc: Any, ph: str) -> list[tuple[Any, Any]]:
    out: list[tuple[Any, Any]] = []
    for page in doc:
        for rect in page.search_for(ph):
            out.append((page, rect))
    return out


def _pdf_redact_insert_text(page: Any, rect: Any, val: str) -> None:
    try:
        page.add_redact_annot(
            rect,
            text=val,
            fontsize=10,
            fontname="helv",
            text_color=(0, 0, 0),
            fill=(1, 1, 1),
        )
        page.apply_redactions()
    except TypeError:
        page.add_redact_annot(rect)
        page.apply_redactions()
        page.insert_text(rect.tl, val, fontsize=10, fontname="helv", color=(0, 0, 0))


def fill_pdf(blob: bytes, schema: list[dict[str, Any]], answers: dict[str, Any]) -> bytes:
    try:
        import pymupdf as fitz
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF fill support (PyMuPDF) is not available.",
        ) from exc

    doc = fitz.open(stream=blob, filetype="pdf")
    specs = _placeholder_specs(schema, answers)
    strip_specs = _radio_group_label_bracket_strip_specs(schema)
    pdf_radio_strip = _explicit_radio_markup_strip_specs_from_pdf(doc)

    image_specs = [s for s in specs if isinstance(s[1], bytes)]
    text_specs = strip_specs + pdf_radio_strip + [s for s in specs if isinstance(s[1], str)]

    for ph, img_bytes, midx in image_specs:
        if midx is None:
            for page in doc:
                while True:
                    hits = page.search_for(ph)
                    if not hits:
                        break
                    rect = hits[0]
                    page.add_redact_annot(rect)
                    page.apply_redactions()
                    try:
                        page.insert_image(rect, stream=img_bytes, keep_proportion=True)
                    except Exception:
                        page.insert_image(rect, stream=img_bytes)
        else:
            hits = _pdf_collect_hits(doc, ph)
            if midx < len(hits):
                page, rect = hits[midx]
                page.add_redact_annot(rect)
                page.apply_redactions()
                try:
                    page.insert_image(rect, stream=img_bytes, keep_proportion=True)
                except Exception:
                    page.insert_image(rect, stream=img_bytes)

    text_specs_sorted = sorted(
        text_specs,
        key=lambda t: (
            t[0],
            1 if t[2] is None else 0,
            0 if t[2] is None else -int(t[2]),
        ),
    )
    for ph, val, midx in text_specs_sorted:
        if midx is None:
            for page in doc:
                while True:
                    hits = page.search_for(ph)
                    if not hits:
                        break
                    rect = hits[0]
                    _pdf_redact_insert_text(page, rect, val)
        else:
            hits = _pdf_collect_hits(doc, ph)
            if midx < len(hits):
                page, rect = hits[midx]
                _pdf_redact_insert_text(page, rect, val)

    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def fill_image_overlay(blob: bytes, schema: list[dict[str, Any]], answers: dict[str, Any]) -> bytes:
    img = Image.open(io.BytesIO(blob)).convert("RGB")
    w, h = img.size
    try:
        font = ImageFont.truetype("arial.ttf", 15)
    except OSError:
        font = ImageFont.load_default()

    row_blocks: list[tuple[str, Image.Image | None]] = []
    extra_h = 24
    for row in schema:
        key = row["key"]
        raw = str(answers.get(key, "") if answers.get(key) is not None else "")
        it = str(row.get("input_type") or "text").strip().lower()
        label = str(row.get("label") or key)
        if it == "signature":
            ib = _image_bytes_for_embed(raw)
            if ib:
                try:
                    sig = Image.open(io.BytesIO(ib)).convert("RGBA")
                    sig.thumbnail((max(120, w - 36), 100))
                    row_blocks.append((label, sig))
                    extra_h += sig.size[1] + 28
                    continue
                except Exception:
                    pass
            row_blocks.append((f"{label}: {_signature_fallback_text(raw)}", None))
            extra_h += 22
        elif it == "checkbox":
            row_blocks.append((f"{label}: {_checkbox_merge_replacement(answers.get(key))}", None))
            extra_h += 22
        elif it == "radio":
            rg = str(row.get("radio_group") or "").strip()
            gk = rg if rg else key
            sel = str(answers.get(gk, "") or "").strip()
            opt = str(row.get("radio_option") or key).strip()
            row_blocks.append(
                (f"{label}: {_checkbox_merge_replacement(bool(sel) and sel.lower() == opt.lower())}", None)
            )
            extra_h += 22
        else:
            row_blocks.append((f"{label}: {raw}", None))
            extra_h += 22

    extra = min(640, max(100, extra_h))
    new_img = Image.new("RGB", (w, h + extra), (252, 252, 254))
    new_img.paste(img, (0, 0))
    draw = ImageDraw.Draw(new_img)
    y = h + 12
    for text_line, sig_img in row_blocks:
        if sig_img is None:
            draw.text((14, y), text_line, fill=(30, 27, 45), font=font, spacing=4)
            y += 22
        else:
            draw.text((14, y), f"{text_line}", fill=(30, 27, 45), font=font, spacing=4)
            y += 20
            new_img.paste(sig_img, (14, y), sig_img)
            y += sig_img.size[1] + 10

    buf = io.BytesIO()
    new_img.save(buf, format="PNG")
    return buf.getvalue()
